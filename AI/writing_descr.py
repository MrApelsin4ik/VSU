import os
import re
import json
import time
import logging
import argparse
import tempfile
import shutil
from typing import List, Dict, Any, Optional
from urllib.parse import urlparse, unquote
import mimetypes

import requests
from dotenv import load_dotenv


try:
    load_dotenv(".env")
except Exception:
    pass

# ----------------------- конфигурация -----------------------
API_KEY = os.getenv("API_KEY")
API_URL = os.getenv("API_URL", "http://127.0.0.1:9000").rstrip("/")
LM_API_URL = os.getenv("LM_API_URL", "http://192.168.0.117:1234/v1/chat/completions").rstrip("/")
LM_API_KEY = os.getenv("LM_API_KEY", None)
LM_MODEL = os.getenv("LM_MODEL", 'openai/gpt-oss-20b')

# параметры chunking / summarization
CHUNK_SIZE_CHARS = int(os.getenv("CHUNK_SIZE_CHARS", "3000"))
CHUNK_OVERLAP_CHARS = int(os.getenv("CHUNK_OVERLAP_CHARS", "400"))
MAX_RETRIES = 3
SLEEP_BETWEEN_REQUESTS = float(os.getenv("SLEEP_BETWEEN_REQUESTS", "0.2"))

# логирование
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

# ----------------------- HTTP helpers -----------------------
TIMEOUT = 120.0


def _headers(api_key: Optional[str] = None) -> Dict[str, str]:
    h = {"Content-Type": "application/json"}
    key = api_key if api_key is not None else API_KEY
    if key:
        h["X-API-KEY"] = key
    return h


class SimpleAPIClient:
    def __init__(self, base_url: str = API_URL, api_key: Optional[str] = API_KEY):
        self.base_url = base_url.rstrip("/")
        self.api_key = api_key

    def _url(self, path: str) -> str:
        return f"{self.base_url}{path}"

    def get_missing(self) -> Dict[str, Any]:
        url = self._url("/api/ai/missing/")
        r = requests.get(url, headers=_headers(self.api_key), timeout=TIMEOUT)
        r.raise_for_status()
        return r.json()

    def get_full(self, model: str, pk: int) -> Dict[str, Any]:
        url = self._url(f"/api/ai/full/{model}/{pk}/")
        r = requests.get(url, headers=_headers(self.api_key), timeout=TIMEOUT)
        r.raise_for_status()
        return r.json()

    def update_item(self, model: str, pk: int, description_ai: Optional[str] = None, description_ai_l2: Optional[str] = None) -> Dict[str, Any]:
        data = {"model": model, "id": pk}
        if description_ai is not None:
            data["description_ai"] = description_ai
        if description_ai_l2 is not None:
            data["description_ai_l2"] = description_ai_l2
        url = self._url("/api/ai/update/")
        r = requests.post(url, headers=_headers(self.api_key), data=json.dumps(data), timeout=TIMEOUT)
        r.raise_for_status()
        return r.json()


# ----------------------- текстовые утилиты -----------------------
HTML_TAG_RE = re.compile(r"<[^>]+>")
WHITESPACE_RE = re.compile(r"\s+")


def strip_html(html: str) -> str:
    if not html:
        return ""
    text = HTML_TAG_RE.sub(" ", html)
    text = WHITESPACE_RE.sub(" ", text).strip()
    return text


def extract_best_text(payload: Dict[str, Any], tmp_dir: Optional[str] = None) -> str:
    """
    Пытается найти наиболее подходящее текстовое поле в возвращаемом API get_full.
    Также скачивает и извлекает текст из вложений (pdf/docx/txt) и добавляет в конец текста.
    """
    candidates = []
    for key in ("text", "content", "body", "description", "html", "full_text", "article"):
        v = payload.get(key)
        if v and isinstance(v, str) and v.strip():
            if key == "html":
                candidates.append(strip_html(v))
            else:
                candidates.append(v.strip())


    if not candidates:
        if isinstance(payload.get("blocks"), list):
            parts = []
            for b in payload["blocks"]:
                if isinstance(b, dict):
                    for k in ("text", "content"):
                        if b.get(k):
                            parts.append(b.get(k))
                elif isinstance(b, str):
                    parts.append(b)
            if parts:
                candidates.append(" ".join(parts))


    attachment_texts = []
    att = payload.get("attachments") or payload.get("files") or payload.get("resources")
    if att and isinstance(att, list):
        # create tmp dir if needed
        if tmp_dir is None:
            tmp_dir = tempfile.mkdtemp(prefix="ingest_")
        else:
            os.makedirs(tmp_dir, exist_ok=True)
        for it in att:
            url = None
            filename_hint = None
            if isinstance(it, dict):
                url = it.get("url") or it.get("file") or it.get("download_url")
                filename_hint = it.get("filename") or it.get("name") or it.get("title")
            elif isinstance(it, str):
                url = it
            if not url:
                continue
            try:
                local = download_file(url, tmp_dir, filename_hint=filename_hint)
                text_from_file = extract_text_from_file(local)
                if text_from_file:
                    attachment_texts.append(f"[file: {os.path.basename(local)}] " + text_from_file[:15000])
            except Exception as e:
                logging.warning("Failed to download/extract attachment %s: %s", url, e)

        if tmp_dir and os.path.exists(tmp_dir):
            try:
                shutil.rmtree(tmp_dir)
            except Exception:
                pass

    if attachment_texts:
        candidates.append("\n\n".join(attachment_texts))

    if not candidates:

        text = json.dumps(payload, ensure_ascii=False)
        return text[:10000]
    candidates = sorted(candidates, key=lambda s: len(s), reverse=True)
    return candidates[0]


import re
from typing import List

SENTENCE_END_RE = re.compile(r'[.!?](?:\s+|$)')

def chunk_text(text: str, size: int = 500, overlap: int = 50) -> List[str]:
    if not text:
        return []

    chunks = []
    n = len(text)
    i = 0

    while i < n:
        raw_end = min(i + size, n)
        chunk = text[i:raw_end]

        ends = list(SENTENCE_END_RE.finditer(chunk))

        if ends:
            last_end_pos = ends[-1].end()
            end = i + last_end_pos
        else:

            space_pos = chunk.rfind(' ')
            end = i + (space_pos if space_pos != -1 else raw_end - i)

        chunk_text_part = text[i:end].strip()
        if chunk_text_part:
            chunks.append(chunk_text_part)

        if end >= n:
            break

        i = max(i, end - overlap)

    return chunks



# ----------------------- скачивание и обработка файлов -----------------------
def download_file(url: str, dest_dir: str, filename_hint: Optional[str] = None) -> str:
    """
    Скачивает файл по URL в dest_dir и возвращает путь к локальному файлу.
    Работает с HTTP/HTTPS.
    """
    parsed = urlparse(url)
    if not parsed.scheme.startswith("http"):
        raise ValueError(f"Unsupported URL scheme: {url}")

    name = None
    if filename_hint:
        name = filename_hint
    else:
        name = os.path.basename(unquote(parsed.path)) or None
    if not name:
        name = f"file_{int(time.time()*1000)}"
    local_path = os.path.join(dest_dir, name)

    with requests.get(url, stream=True, timeout=TIMEOUT) as r:
        r.raise_for_status()
        with open(local_path, "wb") as f:
            for chunk in r.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
    return local_path


def extract_text_from_file(path: str) -> str:
    """
    По расширению вызывает соответствующий извлекатель.
    Поддерживаются: .pdf, .docx, .txt
    """
    lower = path.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(path)
    if lower.endswith(".docx"):
        return extract_text_from_docx(path)
    if lower.endswith(".txt"):
        return extract_text_from_txt(path)
    if lower.endswith(".jpg") or lower.endswith(".jpeg") or lower.endswith(".png"):

        logging.warning("Ignored image file: %s", path)
        return ""

    mt, _ = mimetypes.guess_type(path)
    if mt == "application/pdf":
        return extract_text_from_pdf(path)

    try:
        return extract_text_from_txt(path)
    except Exception:
        return ""


def extract_text_from_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:

        try:
            with open(path, "r", encoding="cp1251") as f:
                return f.read()
        except Exception:
            logging.exception("Failed to read txt file %s", path)
            return ""
    except Exception:
        return ""


def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document
    except Exception as e:
        logging.error("python-docx is required for docx extraction. Install via 'pip install python-docx'")
        raise
    try:
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        # also try tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception:
        logging.exception("Failed to extract docx %s", path)
        return ""


def extract_text_from_pdf(path: str) -> str:
    try:
        import fitz  # PyMuPDF
    except Exception as e:
        logging.error("pymupdf is required for pdf extraction. Install via 'pip install pymupdf'")
        raise
    texts = []
    try:
        doc = fitz.open(path)
        for page in doc:
            try:
                txt = page.get_text("text")
                if txt:
                    texts.append(txt)
            except Exception:

                continue
        doc.close()
    except Exception:
        logging.exception("Failed to extract text from pdf %s", path)
    return "\n".join(texts)


# ----------------------- LLM-инференс -----------------------
def extract_final_fragment(raw_text: str) -> str:
    """
    Извлекает финальную часть ответа LLM.
    Поддерживает несколько форматов:
      - маркеры вида <|final|> или <|channel|>final<|message|>...
      - блоки вида <|start|>assistant<|channel|>final<|message|>...<|end|>
      - блоки ```json ... ```
      - JSON-подстроки {...}
      - fallback — последние непустые строки
    Возвращает пустую строку, если ничего не найдено.
    """
    if not raw_text:
        return ""

    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")

    m = re.search(r"<\|channel\|>final<\|message\|>(.*)", text, flags=re.S)
    if m:
        return m.group(1).strip()

    m = re.search(r"<\|start\|\s*assistant\s*\|>\s*<\|channel\|\s*final\s*\|>\s*<\|message\|\s*(.*?)\s*<\|end\|>", text, flags=re.S|re.I)
    if m:
        return m.group(1).strip()


    m = re.search(r"<\|channel\|\s*final\s*\|>(.*?)(?:<\|channel\||<\|end\||$)", text, flags=re.S|re.I)
    if m:
        return m.group(1).strip()


    m = re.search(r"<\|final\|>(.*)$", text, flags=re.S|re.I)
    if m:
        return m.group(1).strip()


    m = re.search(r"(?mi)^[\s\#\-]*final\s*[:\-]*\s*$", text)
    if m:
        start = m.end()
        return text[start:].strip()


    m = re.search(r"(?i)final\s*[:\-]\s*(.*)$", text, flags=re.S|re.M)
    if m:
        return m.group(1).strip()


    m = re.search(r"```json(.*?)```", text, flags=re.S|re.I)
    if m:
        return m.group(1).strip()


    m = re.search(r"<<JSON>>(.*?)<</JSON>>", text, flags=re.S)
    if m:
        return m.group(1).strip()


    json_candidates = list(re.finditer(r"\{[\s\S]*?\}", text))
    for jm in reversed(json_candidates):
        candidate = jm.group(0)
        try:
            json.loads(candidate)
            return candidate.strip()
        except Exception:
            cand2 = candidate.replace("'", '"')
            try:
                json.loads(cand2)
                return cand2.strip()
            except Exception:
                continue


    m = re.search(r"<\|start\|\s*assistant.*?\|>(.*)$", text, flags=re.S|re.I)
    if m:
        tail = m.group(1).strip()

        m2 = re.search(r"<\|channel\|\s*final\s*\|>(.*)$", tail, flags=re.S|re.I)
        if m2:
            return m2.group(1).strip()
        return tail


    lines = [ln.rstrip() for ln in text.splitlines() if ln.strip()]
    if lines:
        tail = "\n".join(lines[-3:])
        return tail.strip()

    return text.strip()


def _clean_llm_text(s: str) -> str:
    if not s:
        return s


    s = re.sub(r"<\|[^|]+\|>", " ", s)


    s = re.sub(r"[\x00-\x1f\x7f]+", " ", s)

    s = WHITESPACE_RE.sub(" ", s).strip()
    return s

def call_llm(prompt: str, system_prompt: Optional[str] = None, max_tokens: int = 512, temperature: float = 0.0) -> str:
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": prompt})

    payload = {
        "model": LM_MODEL,
        "messages": messages,
        "temperature": temperature,
    }
    print('\n\n\n')
    print(payload)
    headers = {"Content-Type": "application/json"}
    if LM_API_KEY:
        headers["Authorization"] = f"Bearer {LM_API_KEY}"

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            r = requests.post(LM_API_URL, headers=headers, json=payload, timeout=TIMEOUT)
            r.raise_for_status()
            data = r.json()
            print(data)
            print('\n\n\n')

            raw_text = ""
            if isinstance(data, dict):

                if "choices" in data and isinstance(data["choices"], list) and data["choices"]:
                    first = data["choices"][0]

                    if isinstance(first.get("message"), dict):
                        cont = first["message"].get("content") or first["message"].get("text") or ""
                        if isinstance(cont, str) and cont.strip():
                            raw_text = cont

                    if not raw_text and isinstance(first.get("text"), str):
                        raw_text = first["text"]

                    if not raw_text:
                        for k in ("content", "text", "message"):
                            if isinstance(first.get(k), str) and first.get(k).strip():
                                raw_text = first.get(k)
                                break


                if not raw_text:
                    for k in ("text", "result"):
                        if isinstance(data.get(k), str) and data.get(k).strip():
                            raw_text = data.get(k)
                            break

                if not raw_text and "outputs" in data and isinstance(data["outputs"], list) and data["outputs"]:
                    first = data["outputs"][0]
                    if isinstance(first, dict):
                        for k in ("text", "content", "generated_text"):
                            if isinstance(first.get(k), str) and first.get(k).strip():
                                raw_text = first.get(k)
                                break

                if not raw_text and "data" in data and isinstance(data["data"], list) and data["data"]:
                    d0 = data["data"][0]
                    if isinstance(d0, dict):
                        if isinstance(d0.get("generated_text"), str) and d0.get("generated_text").strip():
                            raw_text = d0.get("generated_text")

            if not raw_text:
                raw_text = json.dumps(data, ensure_ascii=False)


            raw_text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]+", " ", raw_text)


            final_fragment = extract_final_fragment(raw_text)
            if final_fragment:
                return _clean_llm_text(final_fragment)

            return _clean_llm_text(raw_text)
        except Exception as e:
            logging.warning("LLM request failed (attempt %d/%d): %s", attempt, MAX_RETRIES, e)
            time.sleep(1.0 * attempt)
    raise RuntimeError("LLM inference failed after retries")


# ----------------------- промпты -----------------------
SYSTEM_PROMPT = (
    "Ты — ассистент, который делает короткие и расширенные описания для контента сайта. "
    "Отвечай на русском. Давай аккуратно и ёмко."
)

PROMPT_CHUNK_SUMMARY = (
    "Сделай очень короткое резюме этого фрагмента (1-2 предложения), затем 3 ключевых факта (по 1-2 коротких предложения). "
    "Если есть явные даты, имена, метки — включи их в ключевые факты. "
    "Фрагмент:\n\n'''{chunk}'''"
)



PROMPT_AGGREGATE_L1 = (
    "У тебя есть набор кратких резюме фрагментов (каждое — 1-2 предложения). "
    "На их основе составь короткое однострочное описание (<= 25 слов) и однопредложный TL;DR (<= 40 слов). "
    "Отвечай простым текстом."
)

PROMPT_AGGREGATE_L2 = (
    "У тебя есть те же краткие резюме фрагментов и исходные длинные фрагменты (если они доступны). "
    "Составь расширённое описание (1-3 коротких абзаца, 120-300 слов) с ключевыми фактами и короткой секцией 'Что важно' (3 буллета). "
    "Отвечай простым текстом."
)


# ----------------------- основной workflow -----------------------
def summarize_document_two_levels(text: str) -> Dict[str, Optional[str]]:
    if not text or not text.strip():
        return {"description_ai": None, "description_ai_l2": None}

    chunks = chunk_text(text, size=CHUNK_SIZE_CHARS, overlap=CHUNK_OVERLAP_CHARS)
    logging.info("Document split to %d chunks (approx %d chars each)", len(chunks), CHUNK_SIZE_CHARS)

    chunk_summaries = []
    for i, ch in enumerate(chunks):

        prompt = PROMPT_CHUNK_SUMMARY.format(chunk=ch)
        print(f'\n\n\nprompt_{i}: {prompt}\n\n\n')
        try:
            resp = call_llm(prompt, system_prompt=SYSTEM_PROMPT, max_tokens=300, temperature=0.0)
            print(f'resp:{resp}')
            logging.debug("Chunk %d resp (raw): %s", i, resp)
        except Exception as e:
            logging.error("LLM failed on chunk %d: %s", i, e)
            resp = ""
        final_resp = extract_final_fragment(resp) or resp
        final_resp = _clean_llm_text(final_resp)
        chunk_summaries.append(final_resp)
        time.sleep(SLEEP_BETWEEN_REQUESTS)

    joined_summaries = "\n---\n".join(chunk_summaries)
    print(f'\n\n\nJOINED_SUMMARIES:\n{joined_summaries}')
    # L1
    prompt_l1 = PROMPT_AGGREGATE_L1 + joined_summaries
    print(f'\n\n{prompt_l1}')
    try:
        out_l1 = call_llm(prompt_l1, system_prompt=SYSTEM_PROMPT, max_tokens=200, temperature=0.0)
        logging.debug("RAW L1 output: %s", out_l1)
    except Exception as e:
        logging.error("LLM failed on aggregate L1: %s", e)
        out_l1 = ""
    print(f'\n\nout_l1: {out_l1}\n\n')
    raw_l1 = out_l1 or ""
    final_fragment_l1 = extract_final_fragment(raw_l1)

    description_ai = None
    if final_fragment_l1:
        cleaned = _clean_llm_text(final_fragment_l1)

        first_line = ""
        for ln in cleaned.splitlines():
            if ln.strip():
                first_line = ln.strip()
                break
        if first_line:
            description_ai = first_line[:400]
        else:
            description_ai = cleaned.splitlines()[0].strip()[:400]

    # L2
    prompt_l2 = PROMPT_AGGREGATE_L2 + joined_summaries
    try:
        out_l2 = call_llm(prompt_l2, system_prompt=SYSTEM_PROMPT, max_tokens=700, temperature=0.0)
        logging.debug("RAW L2 output: %s", out_l2)
    except Exception as e:
        logging.error("LLM failed on aggregate L2: %s", e)
        out_l2 = ""
    print(f'\n\nout_l2: {out_l2}\n\n')
    raw_l2 = out_l2 or ""
    final_fragment_l2 = extract_final_fragment(raw_l2)

    description_ai_l2 = None
    if final_fragment_l2:
        cleaned2 = _clean_llm_text(final_fragment_l2)

        description_ai_l2 = cleaned2.strip()[:1500]

    logging.info("Produced description_ai len=%s, description_ai_l2 len=%s", len(description_ai or ""), len(description_ai_l2 or ""))

    return {"description_ai": description_ai, "description_ai_l2": description_ai_l2}


def process_missing_items(limit: Optional[int] = None, dry_run: bool = False):
    client = SimpleAPIClient()
    res = client.get_missing()
    items = []
    if isinstance(res, dict) and "items" in res:
        items = res["items"]
    elif isinstance(res, list):
        items = res
    elif isinstance(res, dict):
        items = [res]
    else:
        logging.error("Unexpected response from get_missing(): %r", res)
        return

    if limit:
        items = items[:limit]
    logging.info("Found %d missing items to process", len(items))

    for entry in items:
        tmp_dir = None
        try:
            model = entry.get("model") or entry.get("type") or "news"
            pk = entry.get("id") or entry.get("pk") or entry.get("pk_id")
            if pk is None:
                logging.warning("Skipping item without id: %r", entry)
                continue
            logging.info("Processing %s:%s", model, pk)
            full = client.get_full(model, int(pk))

            tmp_dir = tempfile.mkdtemp(prefix="ingest_")
            text = extract_best_text(full, tmp_dir=tmp_dir)
            header = []
            maybe_title = full.get("title") or full.get("name") or full.get("headline")
            if maybe_title:
                header.append(str(maybe_title))
            if full.get("url"):
                header.append(full.get("url"))
            header_str = " — ".join(header) + "\n\n" if header else ""
            out = summarize_document_two_levels(header_str + text)
            logging.info("Result for %s:%s -> l1 len=%s, l2 len=%s", model, pk,
                         len(out.get("description_ai") or ""), len(out.get("description_ai_l2") or ""))
            if dry_run:
                print("DRY RUN — would update:", model, pk)
                print("L1:", out.get("description_ai"))
                print("L2:", out.get("description_ai_l2")[:1000] if out.get("description_ai_l2") else None)
            else:
                update_res = client.update_item(model, int(pk), description_ai=out.get("description_ai"), description_ai_l2=out.get("description_ai_l2"))
                logging.info("Updated %s:%s -> %s", model, pk, update_res)
        except Exception as e:
            logging.exception("Failed processing item %r: %s", entry, e)
        finally:
            if tmp_dir and os.path.exists(tmp_dir):
                try:
                    shutil.rmtree(tmp_dir)
                except Exception:
                    pass
        time.sleep(SLEEP_BETWEEN_REQUESTS)


# ----------------------- CLI -----------------------
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=None, help="Process only first N items")
    parser.add_argument("--dry-run", action="store_true", help="Don't call update_item; just show results")
    args = parser.parse_args()
    logging.info("Starting description generation. LM endpoint: %s", LM_API_URL)
    process_missing_items(limit=args.limit, dry_run=args.dry_run)


if __name__ == "__main__":
    main()
